import re
from docx import Document


def load_document(filepath):
    """Read a .docx file and return all paragraph texts as a list."""
    doc = Document(filepath)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return paragraphs


def find_acronyms(paragraphs):
    """
    Find all acronyms in the text.
    An acronym is 2 or more capital letters, but we skip common
    English words that are sometimes written in caps.
    """
    acronym_pattern = re.compile(r'\b[A-Z]{2,}\b')

    # These are common words that are NOT acronyms
    # We'll ignore these automatically
    common_words_to_skip = {
        'THE', 'AND', 'FOR', 'ARE', 'BUT', 'NOT', 'YOU', 'ALL',
        'CAN', 'HER', 'WAS', 'ONE', 'OUR', 'OUT', 'HAS', 'HAD',
        'HIS', 'HOW', 'ITS', 'WHO', 'DID', 'YES', 'NOW', 'NEW',
        'MAY', 'USE', 'GET', 'GOT', 'LET', 'TWO', 'WAY', 'TOO',
        'ANY', 'SEE', 'SET', 'PUT', 'END', 'ADD', 'OWN', 'BIG'
    }

    found = set()

    for paragraph in paragraphs:
        matches = acronym_pattern.findall(paragraph)
        for match in matches:
            if match not in common_words_to_skip:
                found.add(match)

    return sorted(found)


def confirm_acronyms_with_user(acronyms):
    """
    Show each detected acronym to the user and ask:
    'Is this really an acronym? (y/n)'
    Only keep the ones the user confirms.
    """
    confirmed = []

    print("\n--- Step 1: Confirm Acronyms ---")
    print("I found the following words that look like acronyms.")
    print("Please confirm which ones are actual acronyms.\n")

    for acronym in acronyms:
        answer = input(f"  Is '{acronym}' an acronym? (y/n): ").strip().lower()
        if answer == 'y':
            confirmed.append(acronym)
        else:
            print(f"  OK, skipping '{acronym}'.")

    return confirmed


def get_user_expansions(acronyms):
    """
    For each confirmed acronym, ask the user for the full expansion.
    Returns a dictionary like: {'API': 'Application Programming Interface'}
    """
    expansions = {}

    print("\n--- Step 2: Enter Expansions ---")
    print("Type the full form for each acronym and press Enter.")
    print("If you want to skip one, just press Enter without typing.\n")

    for acronym in acronyms:
        user_input = input(f"  Full form of '{acronym}': ").strip()
        if user_input:
            expansions[acronym] = user_input
        else:
            print(f"  Skipping '{acronym}'.")

    return expansions


def format_glossary(expansions):
    """
    Format the glossary in the style: API(Application Programming Interface)
    """
    glossary = []
    for acronym, expansion in expansions.items():
        glossary.append(f"{acronym}({expansion})")
    return glossary


def main():
    filepath = "data/sample.docx"

    print(f"Reading document: {filepath}")
    paragraphs = load_document(filepath)

    print(f"\nFound {len(paragraphs)} paragraph(s).")
    print("\n--- Document Content ---")
    for i, para in enumerate(paragraphs, 1):
        print(f"  {i}. {para}")

    print("\nSearching for acronyms...")
    acronyms = find_acronyms(paragraphs)

    if not acronyms:
        print("No acronyms found.")
        return

    print(f"\nDetected {len(acronyms)} potential acronym(s): {', '.join(acronyms)}")

    # Step 1: user confirms which ones are real acronyms
    confirmed_acronyms = confirm_acronyms_with_user(acronyms)

    if not confirmed_acronyms:
        print("\nNo acronyms confirmed. Nothing to expand.")
        return

    print(f"\nConfirmed {len(confirmed_acronyms)} acronym(s): {', '.join(confirmed_acronyms)}")

    # Step 2: user provides expansions
    expansions = get_user_expansions(confirmed_acronyms)

    if not expansions:
        print("\nNo expansions provided.")
        return

    # Step 3: format and display the final glossary
    glossary = format_glossary(expansions)

    print("\n--- Final Glossary ---")
    for entry in glossary:
        print(f"  {entry}")

def find_screenshots(filepath):
    """
    Find all paragraphs that look like screenshot placeholders.
    These are bold/italic paragraphs containing [SCREENSHOT: ...]
    Also detects paragraphs just before an image (inline shapes).
    Returns a list of dicts with context information.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument(filepath)

    screenshots = []
    paragraphs = list(doc.paragraphs)

    for i, para in enumerate(paragraphs):
        # Check if paragraph contains a screenshot placeholder
        if '[SCREENSHOT:' in para.text or '[IMAGE:' in para.text:
            # Get surrounding context — paragraph before this one
            context = paragraphs[i-1].text if i > 0 else ""
            screenshots.append({
                "index": i,
                "placeholder": para.text,
                "context": context,
                "type": "screenshot"
            })

    return screenshots


def find_tables(filepath):
    """
    Find all tables in the document.
    For each table, get the paragraph just before it as context.
    Returns a list of dicts with context information.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument(filepath)

    tables_info = []
    paragraphs = list(doc.paragraphs)

    for i, table in enumerate(doc.tables):
        # Get the text from the first row as a preview
        first_row = [cell.text for cell in table.rows[0].cells]
        preview = " | ".join(first_row)

        # Find context — look for paragraphs before the table
        # We use the last paragraph before the table as context
        context = paragraphs[-1].text if paragraphs else ""
        for para in paragraphs:
            if para.text.strip() and "screenshot" not in para.text.lower():
                context = para.text

        tables_info.append({
            "index": i,
            "preview": preview,
            "context": context,
            "type": "table"
        })

    return tables_info

if __name__ == "__main__":
    main()