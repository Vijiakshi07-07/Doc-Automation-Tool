import sys
import os

# This makes sure Python can find files inside the src folder
sys.path.append(os.path.join(os.path.dirname(__file__), "src"))

from acronym_finder import load_document, find_acronyms, confirm_acronyms_with_user, get_user_expansions, format_glossary
from llm_helper import generate_glossary_definition
from database import create_table, save_entry, get_all_entries
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_glossary_to_document(filepath, glossary_entries):
    """
    Takes the original document and adds a glossary table at the end.
    glossary_entries = list of (acronym, expansion, definition) tuples
    """
    doc = Document(filepath)

    # Add a page break before the glossary
    doc.add_page_break()

    # Add glossary heading
    heading = doc.add_heading("Glossary", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add a table with 3 columns: Acronym | Expansion | Definition
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Acronym"
    header_cells[1].text = "Expansion"
    header_cells[2].text = "Definition"

    # Make header text bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add one row per glossary entry
    for acronym, expansion, definition in glossary_entries:
        row_cells = table.add_row().cells
        row_cells[0].text = acronym
        row_cells[1].text = expansion
        row_cells[2].text = definition if definition else ""

    # Save as a new file so we don't overwrite the original
    output_path = filepath.replace(".docx", "_updated.docx")
    doc.save(output_path)
    return output_path


def main():
    filepath = "data/sample.docx"

    print("=" * 55)
    print("   Document Automation Tool")
    print("=" * 55)

    # Step 1: set up the database
    print("\n[1/6] Setting up database...")
    create_table()

    # Step 2: read the document
    print(f"\n[2/6] Reading document: {filepath}")
    paragraphs = load_document(filepath)
    print(f"      Found {len(paragraphs)} paragraph(s).")

    # Step 3: find acronyms
    print("\n[3/6] Scanning for acronyms...")
    acronyms = find_acronyms(paragraphs)

    if not acronyms:
        print("      No acronyms found. Exiting.")
        return

    print(f"      Detected: {', '.join(acronyms)}")

    # Step 4: user confirms which are real acronyms
    print("\n[4/6] Please confirm acronyms:")
    confirmed = confirm_acronyms_with_user(acronyms)

    if not confirmed:
        print("      No acronyms confirmed. Exiting.")
        return

    # Step 5: user provides expansions
    print("\n[5/6] Please provide expansions:")
    expansions = get_user_expansions(confirmed)

    if not expansions:
        print("      No expansions provided. Exiting.")
        return

    # Step 6: AI generates definitions and everything gets saved
    print("\n[6/6] Generating AI definitions and saving to database...")
    glossary_entries = []

    for acronym, expansion in expansions.items():
        print(f"\n      Processing {acronym}...")

        # Ask AI for a definition
        print(f"      Asking AI for definition...")
        definition = generate_glossary_definition(acronym, expansion)
        print(f"      Definition: {definition}")

        # Save to database
        save_entry(acronym, expansion, definition)

        glossary_entries.append((acronym, expansion, definition))

    # Step 7: write glossary into the Word document
    print("\n[7/7] Writing glossary into document...")
    output_path = add_glossary_to_document(filepath, glossary_entries)
    print(f"      Done! Updated document saved as: {output_path}")

    # Final summary
    print("\n" + "=" * 55)
    print("   Final Glossary")
    print("=" * 55)
    for acronym, expansion, definition in glossary_entries:
        print(f"\n  {acronym}({expansion})")
        print(f"  {definition}")

    print("\n" + "=" * 55)
    print("   All done! Check your data/ folder for the updated document.")
    print("=" * 55)


if __name__ == "__main__":
    main()